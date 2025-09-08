import base64
import glob
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from omegaconf import OmegaConf
from datasets import load_dataset

parent_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
sys.path.append(parent_dir)
print(parent_dir)
from plotting_benchmark.vis_generator import read_responses


def decode_image(encoded_image, output_image_file):
    decoded_image = base64.b64decode(encoded_image)

    # Write the decoded image data to a file
    with open(output_image_file, "wb") as image_file:
        image_file.write(decoded_image)


"""
Just a script to generate a docx file to dump generated tasks and plots to verify them.
"""
do_random = False
# suffix = "_probs"
suffix = ""
if do_random and len(suffix) == 0:
    suffix = "_random"

# config_path = "../configs/config.yaml"
config_path = "configs/config.yaml"
config = OmegaConf.load(config_path)
paths = config.paths

dataset_folder = Path(paths.dataset_folder)
results_folder = Path(paths.out_folder)
temp_folder = results_folder / "temp"
os.makedirs(temp_folder, exist_ok=True)
# 自动查找最新的结果文件
bench_files = list(results_folder.glob("benchmark_stat*.jsonl"))
result_files = list(results_folder.glob("results_*.json"))

if not bench_files:
    raise FileNotFoundError("No benchmark_stat*.jsonl files found in output folder")
if not result_files:
    raise FileNotFoundError("No results_*.json files found in output folder")

# 使用最新的文件
bench_file = sorted(bench_files, key=os.path.getmtime)[-1]
response_file = sorted(result_files, key=os.path.getmtime)[-1]

print(f"Using benchmark file: {bench_file}")
print(f"Using results file: {response_file}")

# 加载数据集获取ground truth
print("Loading dataset for ground truth images...")
dataset = load_dataset("JetBrains-Research/PandasPlotBench", split="test")

# 创建ID到数据集项的映射
dataset_dict = {item['id']: item for item in dataset}
print(f"Loaded {len(dataset_dict)} items from dataset")

# 读取结果文件（pandas JSON格式）
import json
with open(response_file, 'r', encoding='utf-8') as f:
    plot_data = json.load(f)

# 将pandas JSON格式转换为标准格式
plot_responses = {}
if 'id' in plot_data:
    # plot_data的结构是列式的：{"id": {"0": 157, "1": 54, ...}, "plots_generated": {"0": [...], "1": [...], ...}}
    id_dict = plot_data['id']
    
    for index_key, actual_id in id_dict.items():
        entry = {'id': actual_id}
        
        # 对于每个字段，提取对应index_key的值
        for field_name, field_values in plot_data.items():
            if field_name != 'id' and isinstance(field_values, dict):
                entry[field_name] = field_values.get(index_key, None)
        
        plot_responses[actual_id] = entry

print(f"Found {len(plot_responses)} responses")
print(f"Sample IDs: {list(plot_responses.keys())[:3]}")

# 调试：检查每个响应的plots_generated结构
for sample_id in list(plot_responses.keys())[:3]:
    sample_response = plot_responses[sample_id]
    pg = sample_response.get('plots_generated', 'missing')
    print(f"ID {sample_id} plots_generated: {type(pg)} - {str(pg)[:50]}...")

temp_image_file = temp_folder / "plot.png"
temp_gt_image_file = temp_folder / "gt_plot.png"

# 使用结果文件中的ID列表
ids = list(plot_responses.keys())

doc = Document()
section = doc.sections[0]
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height

for idx in ids:
    response = plot_responses[idx]
    
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.add_run(f"ID = {idx}\n")
    
    # 从response中获取分数信息
    if 'score_vis' in response:
        vis_score = response['score_vis']
        paragraph.add_run(f"Vis Score = {vis_score}\n")
    if 'score_task' in response:
        task_score = response['score_task']
        paragraph.add_run(f"Task Score = {task_score}\n")
    if 'has_plot' in response:
        has_plot = response['has_plot']
        paragraph.add_run(f"Has Plot = {has_plot}\n")

    # 从数据集中获取ground truth
    gt_available = False
    if idx in dataset_dict:
        dataset_item = dataset_dict[idx]
        if 'plots_gt' in dataset_item and dataset_item['plots_gt']:
            plots_gt_list = dataset_item['plots_gt']
            if isinstance(plots_gt_list, list) and len(plots_gt_list) > 0:
                gt_available = True
                # 使用列表中的第一个图片
                try:
                    decode_image(plots_gt_list[0], temp_gt_image_file)
                    paragraph.add_run("Ground truth loaded from dataset\n")
                    if len(plots_gt_list) > 1:
                        paragraph.add_run(f"Note: {len(plots_gt_list)} GT images available, using first one\n")
                except Exception as e:
                    paragraph.add_run(f"Error decoding ground truth: {str(e)}\n")
                    gt_available = False
    
    if not gt_available:
        paragraph.add_run(f"Warning: No ground truth found for ID {idx}\n")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    cell = table.cell(0, 0)
    cell.text = "Generated"

    # 调试：输出当前响应的plots_generated信息
    print(f"\n=== ID {idx} Debug Info ===")
    if 'plots_generated' in response:
        pg = response['plots_generated']
        print(f"plots_generated type: {type(pg)}")
        print(f"plots_generated value: {str(pg)[:100]}...")
    else:
        print("No plots_generated in response")

    # 检查生成的图像 - 现在每个条目的plots_generated应该是对应索引的列表
    if 'plots_generated' in response:
        plots_generated = response['plots_generated']
        
        if isinstance(plots_generated, list) and len(plots_generated) > 0:
            # plots_generated现在是该ID对应的图片列表
            image_data = plots_generated[0]  # 取列表中的第一个图片
            if image_data and isinstance(image_data, str) and len(image_data) > 100:
                try:
                    decode_image(image_data, temp_image_file)
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(str(temp_image_file), width=Inches(4))
                    print(f"ID {idx}: Successfully loaded image")
                except Exception as e:
                    print(f"ID {idx}: Error decoding image: {str(e)}")
                    cell.text = f"Generated\nError decoding: {str(e)[:50]}"
            else:
                cell.text = "Generated\nInvalid image data"
        elif isinstance(plots_generated, list) and len(plots_generated) == 0:
            cell.text = "Generated\nNo images (empty list)"
        else:
            cell.text = f"Generated\nUnexpected type: {type(plots_generated)}"
    else:
        cell.text = "Generated\nNo plots_generated field"

    cell = table.cell(0, 1)
    cell.text = "Ground truth"
    if gt_available and temp_gt_image_file.exists():
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(str(temp_gt_image_file), width=Inches(4))
    else:
        cell.text = "Ground truth\nNot available"

    doc.add_page_break()

result_file = results_folder / "bench_results.docx"
doc.save(result_file)
print(f"Output saved in {str(result_file)}")
